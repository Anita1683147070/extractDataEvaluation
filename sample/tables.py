# coding = UTF-8

from docx import Document
import os
import re


def isSelectedTable(tb):
    #rowOk = (len(tb.rows) > 14) & (len(tb.rows) < 33)
    rowOk = 1
    colOk = (len(tb.columns) == 3)
    return rowOk & colOk


def isJDFile(doc):
    isJD = False
    isCompleteJDFile = False
    for para in doc.paragraphs:
        JD = re.search('京东', para.text)
        docHead = re.search('执行案号', para.text)  #可能有多个京东文件，但是只有一份有用
        if JD is not None:
            isJD = True
        if docHead is not None:
            isCompleteJDFile = True
        if isJD and isCompleteJDFile:
            return True
    return False


def add_table2doc(tb, doc):
    table = doc.add_table(rows=len(tb.rows), cols=len(tb.columns), style='Table Grid')
    for rowIndex, row in enumerate(tb.rows):
        for cellIndex, cell in enumerate(row.cells):
            table.rows[rowIndex].cells[cellIndex].text = cell.text
    #return doc


tableNum = 0
dataPath = os.path.abspath('../data')
dirIndex = 0
for root, dirs, files in os.walk(dataPath):
    dirIndex += 1
    tableNum = 0
    table_doc = Document()
    for file in files:
        if os.path.splitext(file)[-1] == '.docx':
            # print('当前文件：', str(file))
            doc = Document(os.path.join(root, file))
            if isJDFile(doc):
                for tb in doc.tables:
                    # print('row', len(tb.rows), 'cell', len(tb.rows[0].cells))
                    if isSelectedTable(tb): #凡是3列的表格
                        add_table2doc(tb, table_doc)
                        tableNum += 1
    table_doc.save('./tempTableDoc/'+str(dirIndex)+'tableInADoc.docx')
    print('dirIndex', dirIndex, ': ', tableNum, '个表格')


