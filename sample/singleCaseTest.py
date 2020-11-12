from docx import Document
import os
import re
import numpy as np
import tensorflow as tf
from bert_serving.client import BertClient
import sys
import string
import json

'''
path = os.path.abspath('../data/107678832')
# print(path)
for root, dirs, files in os.walk(path):
    # print(str(root))
    #print(len(files))
    for file in files:
        if os.path.splitext(file)[1] == '.docx':
            #print('当前文件：', str(file))
            doc = Document(os.path.join(root, file))
            for para in doc.paragraphs:
                m = re.search('京东', para.text)
                if m is not None:
                    print(m.group())
            #for tabIndex, tb in enumerate(doc.tables):
                #print('table', tabIndex, 'has', len(tb.rows), 'rows and',len(tb.rows[0].cells),'columns')
'''


# hh = {'11':'我是1', '22':'我是2', '33':'我是3'}

class H:
    def __init__(self):
        self.h1 = 'I am h1'
        self.h2 = []

    def get_h1(self):
        return self.h1

    def set_h1(self, h1):
        self.h1 = h1

    def get_h2(self):
        return self.h2

    def add2h2(self, H2):
        self.h2.append(H2)


class H2:
    def __init__(self):
        self.h2_1 = ''
        self.h2_2 = ''
        self.h2_3 = H3()

    def set_h2_1(self, h2_content):
        self.h2_1 = h2_content

    def set_h2_2(self, h2_content):
        self.h2_2 = h2_content

    def get_h2_1(self):
        return self.h2_1


class H3:
    def __init__(self):
        self.h3_1 = '31'


def convert2json(h):
    return {
        'h1': h.h1,
        'h2': [convertsub2json(h2_) for h2_ in h.h2]
    }


def convertsub2json(H_2):
    return {
        'h2_1': H_2.h2_1,
        'h2_2': H_2.h2_2,
        'h2_3': convertsubsub2json(H_2.h2_3)
    }


def convertsubsub2json(H_3):
    return {
        'h3-1': H_3.h3_1
    }


def convert(testAll):
    return [convert2json(h) for h in testAll]


why = {'hoho': 'hh', 'xixi': 'xx'}

test = []
'''
testH = H()
h_1 = H2()
h_1.set_h2_1('我是h2_1_1')
h_1.set_h2_2('我是h2_1_2')
h_2 = H2()
h_2.set_h2_1('我是h2_2_1')
testH.add2h2(h_1)
testH.add2h2(h_2)
#test.append(testH)
testH2 = H()
test.append(testH2)
print('test', test)
with open('localTestJson.json', 'w+', encoding='utf-8') as jsonFile:
    jsonContent = json.dumps(test, default=convert2json, ensure_ascii=False, indent=4)
    jsonFile.write(jsonContent)
'''
